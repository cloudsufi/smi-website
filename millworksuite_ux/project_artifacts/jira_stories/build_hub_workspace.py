from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SS  = os.path.join(os.path.dirname(__file__), 'screenshots')
OUT = os.path.dirname(__file__)

# ── helpers ────────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color='F0F4F8'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color='1E5BA8'):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.name = 'Calibri'
    return p

def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.name = 'Calibri'
    if color: run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_screenshot(doc, path, caption, width=6.2):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    for run in cap.runs:
        run.font.size = Pt(9); run.italic = True
        run.font.color.rgb = RGBColor(100, 116, 139)

def add_meta_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        shade_cell(c0, 'EBF2FA')
        c0.paragraphs[0].clear(); r = c0.paragraphs[0].add_run(k)
        r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
        c1.paragraphs[0].clear(); r2 = c1.paragraphs[0].add_run(v)
        r2.font.size = Pt(10); r2.font.name = 'Calibri'
    doc.add_paragraph()

def add_data_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], '1E5BA8')
        p = table.rows[0].cells[i].paragraphs[0]; p.clear()
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9.5)
        r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        bg = 'F0F4F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            shade_cell(table.rows[ri+1].cells[ci], bg)
            p = table.rows[ri+1].cells[ci].paragraphs[0]; p.clear()
            r = p.add_run(str(val)); r.font.size = Pt(9.5); r.font.name = 'Calibri'
    doc.add_paragraph()

def add_ac_table(doc, items):
    table = doc.add_table(rows=1+len(items), cols=2)
    table.style = 'Table Grid'
    for cell, txt in zip(table.rows[0].cells, ['#', 'Acceptance Criterion']):
        shade_cell(cell, '1E5BA8')
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(10)
        r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(255,255,255)
    for i, item in enumerate(items, 1):
        row = table.rows[i].cells
        if i % 2 == 0: shade_cell(row[0],'F0F4F8'); shade_cell(row[1],'F0F4F8')
        row[0].paragraphs[0].clear(); r = row[0].paragraphs[0].add_run(str(i))
        r.font.size = Pt(10); r.font.name = 'Calibri'
        row[1].paragraphs[0].clear(); r2 = row[1].paragraphs[0].add_run(item)
        r2.font.size = Pt(10); r2.font.name = 'Calibri'
    doc.add_paragraph()

def add_divider(doc):
    p = doc.add_paragraph('─' * 90)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    for run in p.runs: run.font.color.rgb = RGBColor(200,210,220); run.font.size = Pt(7)

def bul(doc, items):
    for b in items:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'
    doc.add_paragraph()

def set_margins(doc):
    s = doc.sections[0]
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.05); s.right_margin = Cm(3.05)


# ═══════════════════════════════════════════════════════════════════════════
# MWS-005  APP HUB (Home Page)
# ═══════════════════════════════════════════════════════════════════════════
def build_hub():
    doc = Document(); set_margins(doc)

    add_heading(doc, 'MWS-005 — App Hub (Home Page)', level=1)
    add_para(doc, 'Jira Story  ·  MillworkSuite UX Implementation', italic=True, color='445069', size=10)
    add_divider(doc); doc.add_paragraph()

    add_meta_table(doc, [
        ('Story ID',    'MWS-005'),
        ('Epic',        'MWS — App Shell'),
        ('Type',        'Story'),
        ('Priority',    'P0 — Blocker'),
        ('Story Points','8'),
        ('Labels',      'frontend, hub, home, UX'),
        ('Route',       '/  (post-login home)'),
        ('Access',      'All authenticated users'),
    ])

    add_heading(doc, 'User Story', level=2)
    add_para(doc, (
        'As an authenticated MillworkSuite user, I want to see a personalised home dashboard '
        'immediately after signing in so that I can quickly understand my pipeline status, '
        'jump into active work, and navigate to any part of the platform.'
    ), italic=True, size=11)
    doc.add_paragraph()

    add_heading(doc, 'Screen Reference', level=2)
    add_screenshot(doc, f'{SS}/hub.png', 'Fig 1 — App Hub home page (1440 × 900, @2x)')

    add_para(doc, 'Notifications panel (top-right bell icon)', bold=True, size=10, color='1E5BA8')
    add_screenshot(doc, f'{SS}/hub_notifications.png', 'Fig 2 — Notification panel open')

    add_para(doc, 'User chip dropdown (top-right user card)', bold=True, size=10, color='1E5BA8')
    add_screenshot(doc, f'{SS}/hub_user_dropdown.png', 'Fig 3 — User chip dropdown open')
    add_divider(doc)

    # ── Top Bar ──
    add_heading(doc, 'Component 1 — Top Bar (Global Shell)', level=2)
    add_para(doc, 'Persistent across all authenticated screens. Contains:', size=11)
    bul(doc, [
        '• Hamburger menu button (left) — opens the left navigation drawer.',
        '• MillworkSuite logo + wordmark — click navigates to App Hub.',
        '• Help icon button — opens contextual help.',
        '• Bell (notifications) icon — opens the notifications panel (see Component 1b).',
        '• User chip (right) — shows avatar initials, full name, role · company name, and a chevron ↓.',
        '  - Clicking opens a dropdown with: "My Details" (→ Settings) and "Sign Out".',
        '  - Chevron rotates 180° when dropdown is open.',
        '  - Clicking outside the dropdown closes it.',
    ])

    add_para(doc, '1a. Left Navigation Drawer', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• Slides in from the left, starts below the 46 px topbar (does not overlap it).',
        '• Dark overlay covers the rest of the app while drawer is open.',
        '• Clicking anywhere outside the drawer closes it.',
        '• Escape key closes the drawer.',
        '• Contains: Apps group (Estimate expandable with sub-items: Projects, Catalogue, Templates), '
          'Admin group (Users & Roles, Settings).',
        '• Active screen is highlighted with blue background in the nav item.',
    ])

    add_para(doc, '1b. Notifications Panel', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• Slides out from the top-right when bell icon is clicked.',
        '• Bell shows a red dot badge when unread notifications exist.',
        '• Panel lists individual notifications with icon, title, body, and timestamp.',
        '• "View all notifications" button at the bottom.',
        '• Clicking outside the panel closes it.',
    ])

    # ── Greeting ──
    add_heading(doc, 'Component 2 — Greeting Bar', level=2)
    bul(doc, [
        '• Large avatar circle (initials, e.g. "RH") — left side.',
        '• "Good [morning/afternoon/evening], [First name]" — time-aware greeting, bold 18 px.',
        '• Sub-line: "[Day], [Date] · [N] active projects · [N] bids due this week" — dynamic, muted.',
        '• Right side: Search button (⌘K shortcut) and "+ New project" primary button.',
        '• Search button opens global command palette (Cmd/Ctrl + K).',
    ])

    # ── Bid Strip ──
    add_heading(doc, 'Component 3 — 30-Day Bid Deadline Strip', level=2)
    add_para(doc, 'Full-width horizontal timeline spanning the next 30 days, below the greeting bar.', size=11)
    bul(doc, [
        '• Title: "Bid Deadlines · Next 30 Days" with legend: Urgent ≤7d (red), Active (amber), Upcoming (blue), Review (purple).',
        '• Date range label (e.g. "20 May – 19 Jun 2026") — monospace, right-aligned in legend.',
        '• Red shaded zone covers the first 7 days with a dashed right boundary (urgent window).',
        '• Blue "Today" line anchored at the left edge with a "Today" label above.',
        '• Week tick marks at 7-day intervals with date labels (e.g. 27 May, 3 Jun, 10 Jun, 17 Jun).',
        '• Deadline pins: each project appears as a vertical pin (label chip → stem → dot) at its deadline position.',
        '  - Color matches status: red = urgent, amber = active, blue = upcoming, purple = in review.',
        '  - Tall stems (40 px) and short stems (22 px) alternate to reduce label overlap.',
        '  - When 2+ projects share a deadline, a ×N badge appears on a grouped pin.',
        '• Each pin is clickable → navigates to that project\'s workspace.',
        '• Pins have title tooltips on hover.',
    ])

    # ── Hub Body ──
    add_heading(doc, 'Component 4 — Main Column', level=2)

    add_para(doc, '4a. Your Apps Section', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• Section header: "Your Apps" + grid/list view toggle + "Manage apps" link.',
        '• Grid view (default): app cards in a responsive grid.',
        '• Each app card contains: icon (color-coded), app name, description, footer with stats + status tag.',
        '• Estimate card (Active): blue icon, "14 active projects" stat, green "Active" tag. Click → Projects.',
        '• Project card (Beta/Pilot): amber icon, "Pilot · Q3 2026" label, amber "Beta" tag. Dimmed/non-clickable.',
        '• Submittal card (Coming soon): grey icon, "Q4 2026" label, grey "Soon" tag. Dimmed.',
        '• "+ Request an App" placeholder card — outlined dashed style.',
        '• List view: cards collapse to a single horizontal row each.',
    ])

    add_para(doc, '4b. Needs Your Attention Section', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• Section header: "Needs Your Attention" with a red count badge (e.g. "3").',
        '• List of action items, each with: color-coded icon (red/amber/blue), title, meta text, action button.',
        '• Item types: bid deadline warnings (red), AI detection review flags (amber), team comments (blue).',
        '• Clicking the item row or its action button navigates to the relevant workspace.',
    ])

    add_para(doc, '4c. Recent Projects Section', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• Section header: "Recent Projects" + "All projects →" link (→ Projects dashboard).',
        '• Mini table: 4 columns — Project (name + client/GC), Status (tag), Bid Date, Est. Value.',
        '• Shows last 5 projects ordered by most recently accessed.',
        '• Clicking a row navigates to that project\'s workspace.',
        '• Est. Value shown in monospace font, right-aligned.',
    ])

    # ── Sidebar ──
    add_heading(doc, 'Component 5 — Right Sidebar', level=2)

    add_para(doc, '5a. This Week Calendar Widget', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• Header: "This Week" + date range (e.g. "18–24 May 2026").',
        '• 7-day calendar row (Mon–Sun) with day numbers.',
        '• Today is highlighted (blue circle).',
        '• Days with events show a small colored pip below the number.',
        '• Event list below: each event has a color bar, event name, and date/time.',
        '  - Green: regular meetings.',
        '  - Amber: upcoming deadlines.',
        '  - Red: urgent/overdue deadlines (with a warning ⚠ icon).',
    ])

    add_para(doc, '5b. Pipeline Stats Widget', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• 4 stat boxes in a 2×2 grid:',
        '  Open Projects: count + "+N this month" delta (green).',
        '  Win Rate · Last 90d: percentage + "N won of N bid" delta (green/red).',
        '  Due This Week: count + "N urgent" delta (red if >0).',
        '  Pending Reviews: count + "N need you" delta (amber).',
        '• Each box: large bold number, label below, small delta line with colored arrow icon.',
    ])

    add_para(doc, '5c. Team Activity Feed', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• Section header: "Team Activity".',
        '• List of recent team actions, newest first.',
        '• Each item: avatar (color-coded by team member), action text (bold name + action description), timestamp.',
        '• Actions include: created project, uploaded drawings, mapped items, finalized version.',
        '• "You" refers to the signed-in user (shown with default avatar color).',
    ])

    add_divider(doc)

    add_heading(doc, 'Interactions', level=2)
    add_data_table(doc,
        ['Interaction', 'Behaviour'],
        [
            ['Click hamburger (☰)',              'Opens left navigation drawer'],
            ['Click outside drawer',             'Closes drawer'],
            ['Click bell icon',                  'Toggles notifications panel'],
            ['Click user chip',                  'Toggles My Details / Sign Out dropdown'],
            ['Click My Details',                 'Navigates to Settings screen'],
            ['Click Sign Out',                   'Shows confirm dialog → returns to /login'],
            ['Click app card (Estimate)',         'Navigates to Projects Dashboard'],
            ['Click attention item / button',    'Navigates to relevant project workspace'],
            ['Click recent project row',         'Navigates to project workspace'],
            ['Click "All projects →"',           'Navigates to Projects Dashboard'],
            ['Click bid strip pin',              'Navigates to that project\'s workspace'],
            ['Click "+ New project"',            'Opens New Project Wizard'],
            ['Click Search / press ⌘K',         'Opens global command palette / search'],
            ['Click "Manage apps"',              'Placeholder for app management screen'],
        ]
    )

    add_heading(doc, 'Acceptance Criteria', level=2)
    add_ac_table(doc, [
        'Top bar renders with hamburger, logo, help, bell, and user chip on all authenticated screens.',
        'User chip shows avatar, full name, role, and company name (e.g. "Company Admin · SMI Cabinetry").',
        'Clicking user chip opens dropdown with "My Details" and "Sign Out"; chevron rotates 180°.',
        'Clicking outside the chip dropdown closes it.',
        'Clicking "Sign Out" shows a confirmation dialog before navigating to /login.',
        'Left nav drawer slides in below the topbar (46 px from top) and does not overlap it.',
        'Clicking outside the drawer or pressing Escape closes it.',
        'Bell badge shows when unread notifications exist; panel opens/closes correctly.',
        'Greeting shows time-aware salutation, current date, active project count, and bids due count.',
        '"+ New project" button navigates to the New Project Wizard.',
        'Bid deadline strip renders with today line, 7-day urgent zone, week ticks, and all project pins.',
        'Deadline pins are color-coded by status; grouped pins show ×N badge.',
        'Clicking a pin navigates to the correct project workspace.',
        'Your Apps section renders all 3 app cards with correct status tags and click behavior.',
        'Estimate app card stat shows live active project count.',
        'Needs Your Attention section shows items with correct color icons and action buttons.',
        'Attention item count badge matches number of items displayed.',
        'Recent Projects table shows 5 most recent projects with status, bid date, and est. value.',
        'Clicking a recent project row navigates to the workspace.',
        'Pipeline widget shows all 4 stat boxes with correct values and delta lines.',
        'Win Rate box shows percentage + "N won of N bid" sub-line.',
        'This Week calendar highlights today; event pips and event list render correctly.',
        'Team Activity feed shows recent team actions with correct avatars and timestamps.',
        'Page is fully responsive at desktop, tablet, and mobile breakpoints.',
        'Page passes WCAG 2.1 AA audit.',
    ])

    doc.save(f'{OUT}/MWS-005-App-Hub-Home.docx')
    print('✓ MWS-005-App-Hub-Home.docx')


# ═══════════════════════════════════════════════════════════════════════════
# MWS-006  PROJECT WORKSPACE (Estimate Screen)
# ═══════════════════════════════════════════════════════════════════════════
def build_workspace():
    doc = Document(); set_margins(doc)

    add_heading(doc, 'MWS-006 — Project Workspace (Estimate Screen)', level=1)
    add_para(doc, 'Jira Story  ·  MillworkSuite UX Implementation', italic=True, color='445069', size=10)
    add_divider(doc); doc.add_paragraph()

    add_meta_table(doc, [
        ('Story ID',    'MWS-006'),
        ('Epic',        'MWS — Estimate App'),
        ('Type',        'Story'),
        ('Priority',    'P0 — Blocker'),
        ('Story Points','21'),
        ('Labels',      'frontend, estimate, workspace, AI, UX'),
        ('Route',       '/estimate/projects/:id'),
        ('Breadcrumb',  '← Projects  /  [Project Name]'),
    ])

    add_heading(doc, 'User Story', level=2)
    add_para(doc, (
        'As an estimator, I want to open a project and see the AI-detected millwork items overlaid '
        'on the original PDF drawing so that I can review, correct, and price each item, '
        'then export the finalised estimate to Microvellum or Cabinet Vision.'
    ), italic=True, size=11)
    doc.add_paragraph()

    add_heading(doc, 'Screen Reference', level=2)
    add_screenshot(doc, f'{SS}/workspace.png',
                   'Fig 1 — Project Workspace — Estimating mode, A-201 active (1440 × 900, @2x)')
    add_screenshot(doc, f'{SS}/workspace_ai.png',
                   'Fig 2 — Workspace with AI Summary overlay open on the PDF canvas')
    add_divider(doc)

    # ── Project Header ──
    add_heading(doc, 'Component 1 — Project Header Bar', level=2)
    bul(doc, [
        '• "← Projects" back button — navigates back to the Projects Dashboard.',
        '• Project name: "BayCare Manatee — Pediatric Cardiology" — bold H2.',
        '• Status tag: colored pill (e.g. amber "In Progress") — flex-shrink:0, always visible.',
        '• Version indicator: "Ver  Current · v3" with a version icon — shows current revision number.',
        '• Export button (secondary): triggers Microvellum export; shows info toast: "Export started — Generating Microvellum file…"',
        '• Save button (primary blue): saves current state; shows success toast: "BayCare Manatee saved successfully."',
    ])

    # ── Info Strip ──
    add_heading(doc, 'Component 2 — Project Info Strip', level=2)
    add_para(doc, 'A horizontal row of metadata chips below the header. Non-editable (read-only context bar).', size=11)
    add_data_table(doc,
        ['Chip', 'Example Value', 'Highlight'],
        [
            ['Client',    'BayCare Health System',  'Plain text'],
            ['GC',        'Brasfield & Gorrie',      'Plain text'],
            ['Sector',    'Healthcare',              'Plain text'],
            ['Bid Date',  '07 Apr 2026',             'Bold/colored — urgent if near deadline'],
            ['Estimator', 'Rob Hull',                'Plain text'],
            ['Est. Value','$284,420',                'Bold colored (blue)'],
            ['Drafting',  'Microvellum',             'Blue bold — indicates CAD output target'],
        ]
    )

    # ── Mode Bar ──
    add_heading(doc, 'Component 3 — Mode & Tool Bar', level=2)

    add_para(doc, '3a. Mode Tabs', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• Two mode tabs: "Estimating" (active by default) and "Drafting".',
        '• Estimating mode: right panel shows prices, margins, confidence. Focused on pricing.',
        '• Drafting mode: right panel shows CAD export type per item. Focused on production handoff.',
        '• Switching modes changes the right-panel column layout but not the PDF canvas or left rail.',
    ])

    add_para(doc, '3b. Toolbar Tools', bold=True, size=11, color='1E5BA8')
    add_data_table(doc,
        ['Tool', 'Icon', 'Function'],
        [
            ['Select',        'Cursor arrow',      'Default — click to select a detection box on canvas'],
            ['Pan',           'Hand icon',         'Drag to pan the PDF canvas'],
            ['Add Room',      'Rectangle + plus',  'Draw a new room boundary on the canvas'],
            ['Add Elevation', 'Rectangle + line',  'Add a new elevation view marker'],
            ['Measure',       'Ruler/pencil',      'Measure distances on the drawing'],
            ['Markup',        'Pen/arrow',         'Add freehand markup annotations'],
            ['Annotate',      'Note/pen',          'Add text annotations to the drawing'],
            ['AI Re-scan',    'Gear/circle (blue)','Trigger a fresh AI detection pass on the current page'],
        ]
    )

    add_para(doc, '3c. AI Scan Status Pill', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• Shown right-aligned in the modebar: "● AI scanned 14 / 14 sheets · 2:14 ago".',
        '• Green dot indicates scan is complete.',
        '• Amber dot + "Scanning…" shown while a re-scan is in progress.',
        '• Shows how many sheets have been processed and time elapsed since last scan.',
    ])

    # ── 3-Panel Body ──
    add_heading(doc, 'Component 4 — Three-Panel Body Layout', level=2)
    add_para(doc, 'The workspace body is a fixed three-panel layout: Pages rail (left) | PDF Canvas (center) | Estimate Panel (right).', size=11)
    doc.add_paragraph()

    add_para(doc, '4a. Pages Rail (Left, ~230 px)', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• Header: "Pages" + "14 / 14" count in blue monospace.',
        '• Filter tabs: All (14) | Detected (13) | Review (2, amber badge) | Skipped (1).',
        '• Each page item:',
        '  - Thumbnail swatch: blue = detected, amber = needs review, white = not yet processed, grey = skipped.',
        '  - Page number (e.g. A-201) — bold.',
        '  - Page title (e.g. "Reception Elevation") — muted.',
        '  - Status chips: "12 det" (blue), "3 review" (amber), "Skipped" (grey).',
        '• Active page is highlighted with a blue left border and blue background.',
        '• Clicking a page item: loads that sheet into the center canvas and updates the right panel.',
        '• Review filter shows only pages with items needing attention.',
    ])

    add_para(doc, '4b. PDF Canvas (Center)', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• Renders the actual PDF drawing sheet for the selected page.',
        '• Grid overlay (subtle dot grid) behind the PDF.',
        '• Sheet header inside the canvas: page number + title (e.g. "A-201 — RECEPTION ELEVATION") and scale.',
        '• Stamp at the bottom: "DRAWN: RH · 04/02/2026 · SHEET A-201 OF 14".',
        '• AI Detection bounding boxes overlaid on the drawing, color-coded by confidence:',
        '  - Green border + label: high confidence (≥90%) — e.g. "W2D · 98%".',
        '  - Amber border + label + ⚠ warning icon: medium confidence (70–89%) — e.g. "BSB · 79% ⚠".',
        '  - Red border: low confidence (<70%) — needs manual review.',
        '• Each detection box shows the AWI type code + confidence percentage.',
        '• Clicking a detection box highlights the corresponding row in the right panel.',
        '• Canvas zoom controls (bottom right): − | 100% | + | Fit-to-screen.',
        '• Canvas is pannable (drag) and zoomable (scroll wheel).',
    ])

    add_para(doc, '4c. AI Summary Overlay (canvas overlay)', bold=True, size=11, color='1E5BA8')
    bul(doc, [
        '• A floating card appears on the canvas when triggered (via AI Re-scan or auto-open).',
        '• Header: "AI Summary · A-201" with a pulsing spinner icon.',
        '• Body: plain-language summary of all detected items on the page, e.g.:',
        '  "Reception elevation: 2 wall cabinets (W2D, high conf.), 1 sink base (BSB, conf. 79% — review), '
        '  2 full-height tall cabinets (TH), and 1 countertop run (CT-L). Spec: Plam per A-301."',
        '• Close (×) button dismisses the overlay without affecting detections.',
    ])

    add_para(doc, '4d. Estimate Panel (Right)', bold=True, size=11, color='1E5BA8')
    add_para(doc, 'Header: "Estimate Panel" + "Add item" + "Export" buttons. Four tabs:', size=11)
    add_data_table(doc,
        ['Tab', 'Content'],
        [
            ['A-201 (12)',  'Detected items list for the active page — see Item Row detail below'],
            ['Estimate',    'Category breakdown with markup settings (OH&P rate, sales tax) and subtotals by category'],
            ['Rooms (8)',   'Room-level grouping of items — each room shows item count and subtotal'],
            ['CAD Output',  'Per-item CAD export target (Microvellum SKU / Cabinet Vision mapping)'],
        ]
    )

    add_para(doc, 'Item Row (Detections tab) — anatomy:', bold=True, size=11)
    add_data_table(doc,
        ['Element', 'Detail'],
        [
            ['Line number',      '01, 02, 03… — sequential per page'],
            ['Item name',        'Full description + blue "AI" badge (or no badge if manually added)'],
            ['AWI type chip',    'Short code (e.g. W2D, BSB, TH, CT-L) with tooltip on hover'],
            ['Confidence chip',  'Green (≥90%), amber (70–89%), or grey "Manual" for hand-entered items'],
            ['Catalog mapping',  'Green ✓ + SKU code if mapped; amber ⚠ "Check mapping" if gap; catalogue link'],
            ['Price (right)',    'Dollar amount, bold — derived from catalog pricing method'],
            ['Price sub',        'Unit breakdown (e.g. "$28.17/lf", "flat + ADA", "12 lf × $130")'],
            ['Spec line',        'Dimensions + material + notes (e.g. "36" W × 30" H · Plam · soft-close")'],
            ['Flagged row',      '"Review" badge + amber left border — for items with confidence <80% or mapping gaps'],
        ]
    )

    add_para(doc, 'Estimate tab — markup settings:', bold=True, size=11)
    bul(doc, [
        '• OH&P rate: editable input (default 18%), updates all line item prices in real-time.',
        '• Sales tax: editable input (default 0.0%).',
        '• Category breakdown: each product category shows icon, name, item count, sheet reference, and subtotal.',
        '• Categories: Casework, Countertops, Custom Millwork, Wall Panels, etc.',
        '• PROJECT TOTAL shown at the bottom — large bold, sum of all categories with markup applied.',
    ])

    add_divider(doc)

    add_heading(doc, 'Confidence Color System', level=2)
    add_data_table(doc,
        ['Level', 'Threshold', 'Border Color', 'Chip Color', 'Action Required'],
        [
            ['High',   '≥90%',   'Green',  'Green bg',  'None — auto-accepted'],
            ['Medium', '70–89%', 'Amber',  'Amber bg',  'Estimator review recommended'],
            ['Low',    '<70%',   'Red',    'Red bg',    'Manual confirmation required'],
            ['Manual', 'N/A',    'None',   'Grey bg',   'Manually added item, no AI source'],
        ]
    )

    add_heading(doc, 'Export Behaviour', level=2)
    bul(doc, [
        '• Export button in header: triggers export of the full project estimate.',
        '• Export target shown in the info strip ("Drafting: Microvellum").',
        '• In Drafting mode, each item row shows its CAD export SKU/mapping.',
        '• Toast on export start: "Export started — Generating Microvellum file for [Project Name]…"',
        '• Toast on export complete: "Export ready — [filename].mvd ready to download."',
        '• Export tab in the right panel shows per-item CAD Output mapping.',
    ])

    add_heading(doc, 'Versioning', level=2)
    bul(doc, [
        '• Version indicator in the header: "Current · v3".',
        '• Every Save creates a new version automatically if changes were made.',
        '• Won bids are locked (read-only) — no edits possible on a won version.',
        '• Version history is accessible from the version indicator dropdown.',
    ])

    add_heading(doc, 'Acceptance Criteria', level=2)
    add_ac_table(doc, [
        'Project header renders with back button, project name, status tag, version indicator, Export, and Save buttons.',
        'Info strip shows all 7 metadata chips: Client, GC, Sector, Bid Date, Estimator, Est. Value, Drafting target.',
        'Mode bar renders with Estimating and Drafting tabs; toggling changes right-panel layout.',
        'All 8 toolbar tools render with correct icons and tooltips.',
        'AI scan status pill shows scanned sheet count and elapsed time.',
        'Pages rail renders with All/Detected/Review/Skipped filter tabs and correct counts.',
        'Active page is highlighted in the rail; clicking a page loads it in the canvas.',
        'Page items show thumbnail swatch (blue/amber/white/grey), page number, title, and status chips.',
        'PDF canvas renders the active sheet with grid overlay, sheet header, and scale.',
        'AI detection bounding boxes render on the canvas with correct AWI codes and confidence percentages.',
        'Detection box colors match confidence levels: green ≥90%, amber 70–89%.',
        'Clicking a detection box highlights the corresponding row in the right panel.',
        'Canvas zoom controls (−, %, +, Fit) work correctly.',
        'Canvas is pannable by dragging and zoomable by scroll wheel.',
        'AI Summary overlay renders on canvas with plain-language summary and a close button.',
        'Right panel Detections tab lists all items with line number, name, AI badge, AWI code, confidence chip, mapping chip, price, and spec line.',
        'Flagged items (confidence <80% or mapping gap) show amber "Review" badge and amber left border.',
        'Manual items show grey "Manual" chip instead of confidence percentage.',
        'Estimate tab shows markup inputs (OH&P, sales tax) and category breakdown with subtotals.',
        'Changing OH&P rate updates all line item prices in real-time.',
        'PROJECT TOTAL updates whenever items or markup rates change.',
        'Export button triggers export and shows correct info toast.',
        'Save button saves current state and shows success toast.',
        'Page is fully responsive at desktop and tablet breakpoints.',
        'Page passes WCAG 2.1 AA audit.',
    ])

    doc.save(f'{OUT}/MWS-006-Project-Workspace.docx')
    print('✓ MWS-006-Project-Workspace.docx')


build_hub()
build_workspace()
print('\nAll done.')
