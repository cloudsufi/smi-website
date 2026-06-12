from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SS = os.path.join(os.path.dirname(__file__), 'screenshots')
OUT = os.path.dirname(__file__)

# ── helpers ──────────────────────────────────────────────────────────────────

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def shade_cell(cell, hex_color='F0F4F8'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color='1E5BA8'):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.name = 'Calibri'
    return p

def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_screenshot(doc, path, caption, width=6.2):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    for run in cap.runs:
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(100, 116, 139)

def add_meta_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        shade_cell(c0, 'EBF2FA')
        c0.paragraphs[0].clear()
        r = c0.paragraphs[0].add_run(k)
        r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
        c1.paragraphs[0].clear()
        r2 = c1.paragraphs[0].add_run(v)
        r2.font.size = Pt(10); r2.font.name = 'Calibri'
    doc.add_paragraph()

def add_ac_table(doc, items):
    table = doc.add_table(rows=1 + len(items), cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    shade_cell(hdr[0], '1E5BA8')
    shade_cell(hdr[1], '1E5BA8')
    for cell, txt in zip(hdr, ['#', 'Acceptance Criterion']):
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(txt)
        r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, item in enumerate(items, 1):
        row = table.rows[i].cells
        if i % 2 == 0:
            shade_cell(row[0], 'F0F4F8'); shade_cell(row[1], 'F0F4F8')
        row[0].paragraphs[0].clear()
        r = row[0].paragraphs[0].add_run(str(i))
        r.font.size = Pt(10); r.font.name = 'Calibri'
        row[1].paragraphs[0].clear()
        r2 = row[1].paragraphs[0].add_run(item)
        r2.font.size = Pt(10); r2.font.name = 'Calibri'
    doc.add_paragraph()

def add_data_table(doc, headers, rows, col_shades=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], '1E5BA8')
        p = hdr_cells[i].paragraphs[0]
        p.clear()
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9.5); r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = 'F0F4F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            shade_cell(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            p.clear()
            r = p.add_run(str(val))
            r.font.size = Pt(9.5); r.font.name = 'Calibri'
    doc.add_paragraph()

def add_divider(doc):
    p = doc.add_paragraph('─' * 90)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.color.rgb = RGBColor(200, 210, 220)
        run.font.size = Pt(7)

def set_page_margins(doc, top=1, bottom=1, left=1.2, right=1.2):
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.top_margin = Cm(top * 2.54)
    section.bottom_margin = Cm(bottom * 2.54)
    section.left_margin = Cm(left * 2.54)
    section.right_margin = Cm(right * 2.54)


# ═══════════════════════════════════════════════════════════════════════════
# MWS-001  LOGIN PAGE
# ═══════════════════════════════════════════════════════════════════════════
def build_login():
    doc = Document()
    set_page_margins(doc)

    # Cover bar
    add_heading(doc, 'MWS-001 — Login Page', level=1)
    add_para(doc, 'Jira Story  ·  MillworkSuite UX Implementation', italic=True, color='445069', size=10)
    add_divider(doc)
    doc.add_paragraph()

    # Meta
    add_meta_table(doc, [
        ('Story ID',       'MWS-001'),
        ('Epic',           'MWS — Public & Auth Pages'),
        ('Type',           'Story'),
        ('Priority',       'P0 — Blocker'),
        ('Story Points',   '5'),
        ('Labels',         'frontend, auth, UX'),
        ('Route',          '/login'),
    ])

    # User story
    add_heading(doc, 'User Story', level=2)
    add_para(doc, (
        'As a MillworkSuite user, I want to sign in to the platform with my email and password '
        'so that I can access my company\'s estimating workspace securely.'
    ), italic=True, size=11)
    doc.add_paragraph()

    # Screenshot
    add_heading(doc, 'Screen Reference', level=2)
    add_screenshot(doc, f'{SS}/login.png', 'Fig 1 — Login screen (1440 × 900, @2x)')
    add_divider(doc)

    # Components
    add_heading(doc, 'Page Components', level=2)

    sections = [
        ('1. Background', [
            '• Full-viewport dark gradient (deep navy/forest) — no topbar or navigation visible.',
            '• Background uses a millwork-themed photo overlay at reduced opacity.',
        ]),
        ('2. Login Card', [
            '• Centered white card, fixed width ≈ 380 px, rounded corners, soft box-shadow.',
            '• Vertically centered in the viewport.',
        ]),
        ('2a. Logo', [
            '• MillworkSuite SVG logo at height 30 px, top of card.',
        ]),
        ('2b. Heading & Sub-heading', [
            '• H1: "Welcome back" — bold, ~20 px, dark ink.',
            '• Sub: "Continue where you left off." — muted, secondary ink.',
        ]),
        ('2c. Email Field', [
            '• Label: "Email ID" (micro-label, uppercase).',
            '• Input type: email, autocomplete="email".',
            '• Placeholder: rob@smicabinetry.com.',
            '• Full-width, ~36 px height, blue focus ring (--blue-500).',
        ]),
        ('2d. Password Field', [
            '• Label: "Password" (micro-label, uppercase).',
            '• Input type: password, autocomplete="current-password".',
            '• Show/Hide toggle — eye icon inside right edge of input.',
            '  - Click toggles type between password and text.',
            '  - Icon swaps between open-eye and eye-slash.',
            '  - aria-label="Show or hide password" required.',
        ]),
        ('2e. Sign In Button', [
            '• Label: "Sign in" + right-arrow icon (→).',
            '• Full-width, primary blue, ~44 px height.',
            '• Disabled when either field is empty.',
            '• Loading spinner shown while auth request is in flight.',
            '• On success: triggers EULA modal (first login) or navigates to App Hub.',
        ]),
        ('2f. Forgot Password', [
            '• Text link below button: "Forgot password?"',
            '• Navigates to password reset flow (MWS-005).',
        ]),
    ]
    for title, bullets in sections:
        add_para(doc, title, bold=True, size=11, color='1E5BA8')
        for b in bullets:
            p = doc.add_paragraph(b, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.name = 'Calibri'
        doc.add_paragraph()

    add_divider(doc)

    # States table
    add_heading(doc, 'Behaviour & States', level=2)
    add_data_table(doc,
        ['State', 'Behaviour'],
        [
            ['Empty fields',       'Sign in button disabled; no error until submit attempted'],
            ['Invalid email',      'Inline error: "Please enter a valid email address"'],
            ['Wrong credentials',  'Card-level error banner: "Incorrect email or password."'],
            ['Successful auth',    'EULA modal (first login) or direct to App Hub'],
            ['Forgot password',    'Navigate to password reset screen'],
            ['Network error',      '"Could not connect. Check your connection and try again."'],
        ]
    )

    # Accessibility
    add_heading(doc, 'Accessibility Requirements', level=2)
    for item in [
        'All form fields must have associated <label> elements — not just placeholders.',
        'Password toggle must toggle aria-label between "Show password" and "Hide password".',
        'Sign in button is keyboard-submittable via Enter key from any field.',
        'Error messages announced via aria-live="polite".',
        'Focus order: Logo → Email → Password → Show/Hide → Sign in → Forgot password.',
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'

    doc.add_paragraph()
    add_divider(doc)

    # ACs
    add_heading(doc, 'Acceptance Criteria', level=2)
    add_ac_table(doc, [
        'Background is full-viewport dark gradient with no topbar or nav.',
        'MillworkSuite logo renders at correct size, visible against dark bg.',
        'Email and password fields are correctly labelled and focusable.',
        'Password show/hide toggle switches input type and swaps icon correctly.',
        'Sign in button is disabled when either field is empty.',
        'Successful sign in triggers EULA on first login; bypasses on subsequent logins.',
        'Inline validation errors appear for empty / invalid inputs.',
        'Card-level error appears for wrong credentials.',
        '"Forgot password?" link is present and navigates to reset flow.',
        'Page passes WCAG 2.1 AA keyboard navigation and screen-reader audit.',
        'Responsive layout works on mobile, tablet, and desktop.',
    ])

    doc.save(f'{OUT}/MWS-001-Login-Page.docx')
    print('✓ MWS-001-Login-Page.docx')


# ═══════════════════════════════════════════════════════════════════════════
# MWS-002  LANDING PAGE
# ═══════════════════════════════════════════════════════════════════════════
def build_landing():
    doc = Document()
    set_page_margins(doc)

    add_heading(doc, 'MWS-002 — Marketing Landing Page', level=1)
    add_para(doc, 'Jira Story  ·  MillworkSuite UX Implementation', italic=True, color='445069', size=10)
    add_divider(doc)
    doc.add_paragraph()

    add_meta_table(doc, [
        ('Story ID',       'MWS-002'),
        ('Epic',           'MWS — Public & Auth Pages'),
        ('Type',           'Story'),
        ('Priority',       'P0 — Blocker'),
        ('Story Points',   '8'),
        ('Labels',         'frontend, marketing, UX'),
        ('Route',          '/ (public root)'),
    ])

    add_heading(doc, 'User Story', level=2)
    add_para(doc, (
        'As a prospective millwork company, I want to visit the MillworkSuite landing page so that I can '
        'understand the product, see proof of its value, and take action to sign in or book a demo.'
    ), italic=True, size=11)
    doc.add_paragraph()

    add_heading(doc, 'Full Page Screenshot', level=2)
    add_screenshot(doc, f'{SS}/landing.png', 'Fig 1 — Landing page, above the fold (1440 × 900, @2x)')
    add_divider(doc)

    # Sections
    add_heading(doc, 'Section 1 — Top Navigation Bar', level=2)
    for b in [
        '• Logo: MillworkSuite SVG, height 28 px. Click → scroll to top of landing page.',
        '• Nav links (4): Platform, Pricing, Customers, Resources — plain text, hover underline.',
        '• "Sign in" button: secondary/ghost style → navigates to /login.',
        '• "Book a demo" button: primary blue → opens booking modal (future story).',
        '• Nav is sticky — gains a shadow after scrolling past the hero.',
        '• Mobile: links collapse into a hamburger menu.',
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'
    doc.add_paragraph()

    add_heading(doc, 'Section 2 — Hero', level=2)
    subsections = [
        ('2a. Pilot Badge', [
            '• Small pill: pulsing green dot + "App #1 · Estimate · Now in pilot".',
            '• Green dot has a continuous pulse/glow CSS animation.',
        ]),
        ('2b. Headline (H1)', [
            '• "Stop estimating manually." — plain dark ink.',
            '• "Start closing faster." — brand accent color, italic treatment.',
            '• Font: display weight, ~44–52 px.',
        ]),
        ('2c. Sub-headline', [
            '• One paragraph, ~2 lines of secondary ink text.',
            '• Describes: AI-powered, U.S. architectural millwork, PDF → Microvellum/Cabinet Vision.',
        ]),
        ('2d. CTA Row', [
            '• Primary: "Try it now →" — blue filled, arrow icon. Navigates to /estimate (App Hub).',
            '• Secondary: "Watch 90-second demo" — ghost outline. Opens video modal (future story).',
            '• Side-by-side on desktop, stacked on mobile.',
        ]),
        ('2e. Stats Strip', [
            '• 4 metrics in a horizontal row with thin vertical dividers:',
            '  93% — Faster takeoffs',
            '  11x — More bids per week',
            '  $2.4M — Avg revenue lift / yr',
            '  0 — Re-typing into CAD',
            '• Values: ~28–32 px bold display font. Labels: small muted text below.',
        ]),
    ]
    for title, bullets in subsections:
        add_para(doc, title, bold=True, size=11, color='1E5BA8')
        for b in bullets:
            p = doc.add_paragraph(b, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'
        doc.add_paragraph()

    add_heading(doc, 'Section 3 — App Preview Window', level=2)
    add_para(doc, (
        'A faux browser window showing the app in action. Contains three panels side by side:'
    ), size=11)
    add_data_table(doc,
        ['Panel', 'Contents'],
        [
            ['Browser chrome',      '3 dots (red/yellow/green) + address bar: app.millworksuite.com/projects/baycare-manatee'],
            ['Left rail',           'Header "Active Projects" + 4 project rows (name, status, date). First row is highlighted/active.'],
            ['Center — PDF canvas', 'Blueprint-style PDF background. 3 AI detection boxes overlaid: "2 Door Upper", "1 Door Sink Base", "2 Door Tall". Each box is a blue outline with a label chip.'],
            ['Right rail',          'Header "Detected Items" with blue AI badge. 4 line items (number, name, spec, price). Subtotal row at bottom: Reception · Sub — $4,460.'],
        ]
    )

    add_heading(doc, 'Section 4 — Features Grid (6 cards, 3×2)', level=2)
    add_para(doc, 'Headline: "Everything automated. Nothing missed."', bold=True, size=11)
    add_data_table(doc,
        ['#', 'Title', 'Description'],
        [
            ['1', 'PDF intake → detection',       'Forward bid emails or drag-and-drop. Detection runs in under 90 seconds across the full sheet set.'],
            ['2', 'Estimating & Drafting modes',   'Toggle modes with a click. Estimators see margin; drafters see CAD export type. No app-switching.'],
            ['3', 'Direct-to-Microvellum',          'Detected items map to the Company Product Catalogue and export natively to Microvellum or Cabinet Vision.'],
            ['4', 'Your catalog. Your prices.',     'Linear, sqft, tiered, formula-based, flat + add-ons. Map MWS types to internal SKUs once, reuse forever.'],
            ['5', 'Versioned bids, immutable wins', 'Every revision is its own version. Won bids are locked. 90-day soft-delete recovery on everything.'],
            ['6', 'Granular role permissions',      'Roles tied to features, not labels. Grant consultant access without exposing pricing.'],
        ]
    )
    add_para(doc, 'Each card: SVG icon, H3 title, description text. Hover: translateY(-2px) + increased box-shadow.', italic=True, size=10, color='445069')
    doc.add_paragraph()

    add_heading(doc, 'Section 5 — CTA Band', level=2)
    for b in [
        '• Dark/navy full-width band, white text.',
        '• H2: "Your team estimates in days. We do it in minutes."',
        '• Primary button: "Open the demo workspace" → navigates to App Hub.',
        '• Secondary ghost button: "Book onboarding call" (placeholder).',
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'
    doc.add_paragraph()

    add_heading(doc, 'Section 6 — Footer', level=2)
    for b in [
        '• Single-line footer: left — "© 2026 MillworkSuite · Orlando, Florida".',
        '• Right — "v3.2.1.60.3 · pilot release" in monospace font.',
        '• Minimal — no link columns for pilot phase.',
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'
    doc.add_paragraph()
    add_divider(doc)

    add_heading(doc, 'Interaction Table', level=2)
    add_data_table(doc,
        ['Interaction', 'Behaviour'],
        [
            ['Click "Sign in" (nav)',              'Navigate to /login'],
            ['Click "Try it now →" (hero)',        'Navigate to App Hub (/estimate)'],
            ['Click "Open the demo workspace"',    'Navigate to App Hub (/estimate)'],
            ['Click "Book a demo" / "Book call"',  'Open booking modal (future story)'],
            ['Click "Watch 90-second demo"',       'Open video modal (future story)'],
            ['Hover nav links',                    'Underline + color shift to --blue-700'],
            ['Hover feature cards',                'translateY(-2px) + increased box-shadow'],
            ['Scroll past nav',                    'Nav becomes sticky with subtle drop-shadow'],
        ]
    )

    add_heading(doc, 'Responsive Behaviour', level=2)
    add_data_table(doc,
        ['Breakpoint', 'Behaviour'],
        [
            ['Desktop (≥1200px)', 'Full 3-column feature grid; hero + preview side by side'],
            ['Tablet (768–1199px)', '2-column feature grid; preview stacks below hero'],
            ['Mobile (<768px)', '1-column everything; nav → hamburger; CTA buttons stack'],
        ]
    )

    add_heading(doc, 'Acceptance Criteria', level=2)
    add_ac_table(doc, [
        'Sticky nav renders with logo, 4 nav links, Sign in and Book a demo buttons.',
        'Clicking "Sign in" navigates to the login screen.',
        'Hero headline renders with two-tone styling (plain + accent italic).',
        'Pilot badge animates with a pulsing dot.',
        '4 stats render in a horizontal row with correct values and dividers.',
        '"Try it now" and "Open the demo workspace" both navigate to App Hub.',
        'App preview window renders: faux browser chrome, left rail, PDF canvas with 3 detection boxes, right-rail line items and subtotal.',
        '6 feature cards render in a 3×2 grid on desktop with icon, title, and description.',
        'CTA band renders dark with correct headline and two buttons.',
        'Footer renders with copyright (left) and version string (right, monospaced).',
        'Page is fully responsive across desktop, tablet, and mobile breakpoints.',
        'Nav becomes sticky with shadow after scrolling.',
        'Page passes WCAG 2.1 AA audit.',
    ])

    doc.save(f'{OUT}/MWS-002-Landing-Page.docx')
    print('✓ MWS-002-Landing-Page.docx')


# ═══════════════════════════════════════════════════════════════════════════
# MWS-003  PROJECTS DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════
def build_dashboard():
    doc = Document()
    set_page_margins(doc)

    add_heading(doc, 'MWS-003 — Estimate · Projects Dashboard', level=1)
    add_para(doc, 'Jira Story  ·  MillworkSuite UX Implementation', italic=True, color='445069', size=10)
    add_divider(doc)
    doc.add_paragraph()

    add_meta_table(doc, [
        ('Story ID',       'MWS-003'),
        ('Epic',           'MWS — Estimate App'),
        ('Type',           'Story'),
        ('Priority',       'P0 — Blocker'),
        ('Story Points',   '13'),
        ('Labels',         'frontend, estimate, projects, UX'),
        ('Route',          '/estimate/projects'),
        ('Breadcrumb',     'Home / Estimate'),
    ])

    add_heading(doc, 'User Story', level=2)
    add_para(doc, (
        'As an estimator or company admin, I want to see all my company\'s projects in one place '
        'so that I can quickly find, filter, and open a project to begin or continue estimating work.'
    ), italic=True, size=11)
    doc.add_paragraph()

    # --- Screenshots ---
    add_heading(doc, 'Screen References', level=2)

    add_para(doc, 'LIST VIEW (default)', bold=True, size=10, color='1E5BA8')
    add_screenshot(doc, f'{SS}/dashboard_list.png',
                   'Fig 1 — Projects Dashboard — List view (1440 × 900, @2x)')

    add_para(doc, 'GRID VIEW', bold=True, size=10, color='1E5BA8')
    add_screenshot(doc, f'{SS}/dashboard_grid.png',
                   'Fig 2 — Projects Dashboard — Grid view (1440 × 900, @2x)')

    add_para(doc, 'BOARD (KANBAN) VIEW', bold=True, size=10, color='1E5BA8')
    add_screenshot(doc, f'{SS}/dashboard_board.png',
                   'Fig 3 — Projects Dashboard — Board/Kanban view (1440 × 900, @2x)')

    add_divider(doc)

    # --- Page Header ---
    add_heading(doc, 'Component 1 — Page Header', level=2)
    subsections = [
        ('1a. Breadcrumb', [
            '• "Home" — clickable link, navigates to App Hub (/).',
            '• "/" separator.',
            '• "Estimate" — plain text, current location.',
            '• Font: 11 px, --ink-3 muted color.',
        ]),
        ('1b. Title Row', [
            '• H1: "Projects" — bold, 22 px.',
            '• Sub-line (live-computed): "14 active · 3 due this week · $4.2M in pipeline".',
            '• "+ New project" primary blue button — top right. Opens New Project Wizard.',
        ]),
    ]
    for title, bullets in subsections:
        add_para(doc, title, bold=True, size=11, color='1E5BA8')
        for b in bullets:
            p = doc.add_paragraph(b, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'
        doc.add_paragraph()

    # --- Toolbar ---
    add_heading(doc, 'Component 2 — Toolbar', level=2)

    add_para(doc, '2a. Search Input', bold=True, size=11, color='1E5BA8')
    for b in [
        '• Width: ~280 px. Magnifying glass icon on the left.',
        '• Placeholder: "Search projects, GCs, addresses…"',
        '• Real-time filter (debounced ~300 ms) — matches name, GC, address, status.',
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'
    doc.add_paragraph()

    add_para(doc, '2b. Filter Dropdowns', bold=True, size=11, color='1E5BA8')
    add_data_table(doc,
        ['Filter', 'Default', 'Options'],
        [
            ['Status',   'All',       'Not Started, In Progress, Bid Placed, Won, Lost'],
            ['Owner',    'Anyone',    'List of team members in the company workspace'],
            ['Bid date', 'Last 60d',  'Last 7d, Last 30d, Last 60d, Last 90d, Custom range'],
            ['GC',       'Any',       'Dropdown of all GCs in the company workspace'],
        ]
    )
    for b in [
        '• Active (non-default) filter: --blue-100 background, --blue-700 text.',
        '• Filters combine with AND logic across all active selections.',
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'
    doc.add_paragraph()

    add_para(doc, '2c. View Toggle', bold=True, size=11, color='1E5BA8')
    add_data_table(doc,
        ['Button', 'Icon', 'Default', 'Behaviour'],
        [
            ['Grid',  '4-square',      'No',  'Responsive card grid'],
            ['List',  '3 lines',       'Yes', 'Full-width data table'],
            ['Board', '3 column bars', 'No',  'Kanban board by stage'],
        ]
    )
    add_para(doc, 'Selected view persists in localStorage per user.', italic=True, size=10, color='445069')
    doc.add_paragraph()

    # --- Grid View ---
    add_heading(doc, 'Component 3 — Grid View', level=2)
    add_para(doc, 'Responsive card grid (auto-fill, min card 258 px, 12 px gap).', size=11)
    add_para(doc, 'Each card contains:', bold=True, size=11)

    add_data_table(doc,
        ['Area', 'Element', 'Detail'],
        [
            ['Thumbnail (top ~90 px)', 'Document icon',  'SVG, centered, muted blue'],
            ['Thumbnail',              'Status tag',      'Colored pill with dot — bottom-left of thumb'],
            ['Thumbnail',              'Page count',      'e.g. "14 pages" — bottom-right, monospace'],
            ['Body',                   'Project name',    'Bold 13 px, 2-line max with ellipsis'],
            ['Body',                   'Client · GC',     '"SMI · Bid 04/07/2026", muted 11.5 px'],
            ['Body',                   'Stats row',       'Items / Rooms / Bid — 3 metric columns'],
            ['Footer',                 'Team avatars',    'Stacked overlapping circles (initials), max 4 then +N'],
            ['Footer',                 'Due / Outcome',   'Right-aligned: "Due Fri", "Won 04/15", "Lost 04/02"'],
        ]
    )
    add_para(doc, 'Click anywhere on card → navigate to project workspace.', italic=True, size=10, color='445069')
    add_para(doc, 'Hover: card lifts translateY(-2px) with increased shadow.', italic=True, size=10, color='445069')
    doc.add_paragraph()

    # --- List View ---
    add_heading(doc, 'Component 4 — List View (Default)', level=2)
    add_data_table(doc,
        ['Column', 'Width', 'Notes'],
        [
            ['Status',       '110 px',  'Colored status tag with dot'],
            ['Project',      'flex',    'Project name, bold; entire row is clickable'],
            ['GC',           '140 px',  'General contractor name'],
            ['Bid Date',     '100 px',  'MM/DD/YYYY; sortable'],
            ['Pages',        '60 px',   'Right-aligned number'],
            ['Items',        '60 px',   'Right-aligned; "—" if not started'],
            ['Rooms',        '60 px',   'Right-aligned; "—" if not started'],
            ['Total',        '90 px',   'Blue if active; red strikethrough if lost'],
            ['Team',         '90 px',   'Stacked avatar circles'],
            ['Due/Outcome',  '120 px',  '"Due Fri MM/DD", "Won MM/DD", "Lost MM/DD"'],
            ['Actions',      '36 px',   '⋯ icon button — opens context menu'],
        ]
    )
    add_para(doc, 'Table Header behaviour:', bold=True, size=11)
    for b in [
        '• Sticky positioned — stays visible on scroll.',
        '• Background: gradient --blue-50 → white. Bottom border: 2 px solid --blue-700.',
        '• Sortable columns: Project and Bid Date. Click cycles: asc → desc → default.',
        '• Active sort column text shifts to --blue-700.',
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'
    doc.add_paragraph()

    add_para(doc, 'Status Tag Color System:', bold=True, size=11)
    add_data_table(doc,
        ['Status', 'Dot Color', 'Text Color', 'Background'],
        [
            ['Not Started', '--ink-3 grey', '--ink-2',    '--panel'],
            ['In Progress', '--amber',      '--amber',    '--amber-bg'],
            ['Bid Placed',  '--blue-500',   '--blue-700', '--blue-100'],
            ['Won',         '--green',      '--green',    '--green-bg'],
            ['Lost',        '--red',        '--red',      '--red-bg'],
        ]
    )

    add_para(doc, 'Due / Outcome cell styling:', bold=True, size=11)
    add_data_table(doc,
        ['State', 'Style'],
        [
            ['Due > 3 days',    'Normal --ink-2'],
            ['Due ≤ 2 days',    'Red bold, --red color (urgent)'],
            ['Submitted',       'Normal with date'],
            ['Won',             'Green text + date'],
            ['Lost',            'Red text + date'],
        ]
    )

    # --- Board View ---
    add_heading(doc, 'Component 5 — Board (Kanban) View', level=2)
    add_para(doc, 'Horizontal scroll of 5 stage columns, 230 px wide each, 14 px gap.', size=11)
    add_data_table(doc,
        ['Stage', 'Badge Color', 'Accent Bar'],
        [
            ['Not Started', 'Grey background',   '--border grey'],
            ['In Progress', 'Amber background',  '--amber'],
            ['Bid Placed',  'Blue background',   '--blue-500'],
            ['Won',         'Green background',  '--green'],
            ['Lost',        'Red background',    '--red'],
        ]
    )

    add_para(doc, 'Kanban Card structure:', bold=True, size=11)
    for b in [
        '• Project name — bold 12 px, max 2 lines.',
        '• GC + page count — muted 10.5 px.',
        '• Footer: bid value (left, blue bold) + team avatars (right, 20 px, -4 px overlap).',
        '• Due line — color-coded by urgency (normal / amber / green won / red lost).',
        '• Hover: translateY(-1px) + increased shadow.',
        '• Click: navigate to project workspace.',
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'
    doc.add_paragraph()

    # --- Empty States ---
    add_heading(doc, 'Component 6 — Empty States', level=2)
    add_data_table(doc,
        ['Scenario', 'Message'],
        [
            ['No projects exist',        '"No projects yet. Create your first project." + primary + New project button'],
            ['Search returns nothing',   '"No projects match your search." + Clear filters link'],
            ['All filtered out',         '"No projects match the active filters." + Clear filters link'],
        ]
    )

    # --- Sample Data ---
    add_heading(doc, 'Sample Project Data (Implementation Reference)', level=2)
    add_data_table(doc,
        ['Project', 'GC', 'Status', 'Bid Date', 'Pg', 'Items', 'Rooms', 'Total'],
        [
            ['BayCare Manatee — Pediatric Cardiology', 'Brasfield & Gorrie', 'In Progress',  '04/07/2026', '14', '142', '8',  '$284k'],
            ["St. Joe's Children's — 8th Floor Reno",  'Suffolk',            'Bid Placed',   '04/03/2026', '22', '218', '12', '$412k'],
            ['OH ORMC 5th Floor EP & Cath Lab',        'Brasfield & Gorrie', 'Won',          '04/01/2026', '38', '486', '24', '$1.2M'],
            ['AH Punta Gorda HBED — Reno',             'DPR',                'Not Started',  '03/20/2026', '9',  '—',   '—',  '—'],
            ['Baptist Health MD Anderson',              'Skanska',            'In Progress',  '03/09/2026', '17', '96',  '5',  '$176k'],
            ['Nemours Viera MOB SD Build-out',          'Robins & Morton',    'Lost',         '03/24/2026', '28', '312', '14', '$668k'],
        ]
    )

    # --- Permissions ---
    add_heading(doc, 'Role-Based Access', level=2)
    add_data_table(doc,
        ['Role', 'See Projects', 'Create', 'Delete', 'See Financials'],
        [
            ['Company Admin', '✓ All',        '✓', '✓',        '✓'],
            ['Estimator',     '✓ All',        '✓', 'Own only', '✓'],
            ['Drafter',       '✓ All',        '✗', '✗',        '✗'],
            ['Viewer',        '✓ Read-only',  '✗', '✗',        '✗'],
        ]
    )

    # --- Navigation ---
    add_heading(doc, 'Navigation & Routing', level=2)
    add_data_table(doc,
        ['Action', 'Destination'],
        [
            ['Click "Home" breadcrumb',         'App Hub (/)'],
            ['Click any project',               '/estimate/projects/:id (Workspace)'],
            ['Click "+ New project"',           '/estimate/projects/new (Wizard)'],
            ['⋯ → Duplicate',                   'Confirm modal → duplicate → reload list'],
            ['⋯ → Archive',                     'Confirm modal → remove from active list'],
            ['⋯ → Delete',                      'Warning modal → soft-delete (90-day recovery)'],
        ]
    )
    add_divider(doc)

    # --- ACs ---
    add_heading(doc, 'Acceptance Criteria', level=2)
    add_ac_table(doc, [
        'Breadcrumb "Home / Estimate" renders; "Home" navigates to Hub.',
        'Page H1 reads "Projects" with live-computed summary sub-line.',
        '"+ New project" button navigates to the new project wizard.',
        'Search input filters list in real-time by name, GC, and address.',
        'All 4 filter dropdowns render with correct defaults and option lists.',
        'Active (non-default) filters are visually highlighted (blue background).',
        'Multiple filters combine with AND logic.',
        'Grid, List, and Board view toggle buttons switch the view instantly.',
        'Active view button is highlighted in blue; selection persists in localStorage.',
        'Grid view: all 6 sample cards render with thumbnail, status tag, page count, stats, team avatars, due date.',
        'Clicking a card navigates to the workspace.',
        'List view: sticky table header, all 10 columns, correct data, sort indicators on Project and Bid Date.',
        'Clicking a row navigates to workspace; clicking ⋯ does not.',
        'Board view: 5 columns in correct stage order with color coding, count badges, and accent bars.',
        'Kanban cards show name, GC, pages, bid value, team avatars, and due line.',
        'All 3 empty states render correctly.',
        'Status tag color variants render correctly across all three views.',
        'Role-based visibility: financials hidden for Drafters and Viewers.',
        'Page is fully responsive at desktop, tablet, and mobile breakpoints.',
        'Page passes WCAG 2.1 AA audit.',
    ])

    doc.save(f'{OUT}/MWS-003-Projects-Dashboard.docx')
    print('✓ MWS-003-Projects-Dashboard.docx')


build_login()
build_landing()
build_dashboard()
print('\nAll Word documents generated.')
