from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SS  = os.path.join(os.path.dirname(__file__), 'screenshots')
OUT = os.path.dirname(__file__)

# ── shared helpers (same as build_docs.py) ────────────────────────────────

def shade_cell(cell, hex_color='F0F4F8'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color='1E5BA8'):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.name = 'Calibri'
    return p

def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_screenshot(doc, path, caption, width=6.2):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    for run in cap.runs:
        run.font.size = Pt(9); run.italic = True
        run.font.color.rgb = RGBColor(100, 116, 139)

def add_meta_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        shade_cell(c0, 'EBF2FA')
        c0.paragraphs[0].clear(); r = c0.paragraphs[0].add_run(k)
        r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
        c1.paragraphs[0].clear(); r2 = c1.paragraphs[0].add_run(v)
        r2.font.size = Pt(10); r2.font.name = 'Calibri'
    doc.add_paragraph()

def add_data_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], '1E5BA8')
        p = hdr_cells[i].paragraphs[0]; p.clear()
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9.5)
        r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = 'F0F4F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            shade_cell(cells[ci], bg)
            p = cells[ci].paragraphs[0]; p.clear()
            r = p.add_run(str(val)); r.font.size = Pt(9.5); r.font.name = 'Calibri'
    doc.add_paragraph()

def add_ac_table(doc, items):
    table = doc.add_table(rows=1 + len(items), cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ['#', 'Acceptance Criterion']):
        shade_cell(cell, '1E5BA8')
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(10)
        r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(255, 255, 255)
    for i, item in enumerate(items, 1):
        row = table.rows[i].cells
        if i % 2 == 0:
            shade_cell(row[0], 'F0F4F8'); shade_cell(row[1], 'F0F4F8')
        row[0].paragraphs[0].clear(); r = row[0].paragraphs[0].add_run(str(i))
        r.font.size = Pt(10); r.font.name = 'Calibri'
        row[1].paragraphs[0].clear(); r2 = row[1].paragraphs[0].add_run(item)
        r2.font.size = Pt(10); r2.font.name = 'Calibri'
    doc.add_paragraph()

def add_divider(doc):
    p = doc.add_paragraph('─' * 90)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.color.rgb = RGBColor(200, 210, 220); run.font.size = Pt(7)

def bullets(doc, items):
    for b in items:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs: run.font.size = Pt(10.5); run.font.name = 'Calibri'
    doc.add_paragraph()

def set_margins(doc):
    s = doc.sections[0]
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.05); s.right_margin = Cm(3.05)

# ═══════════════════════════════════════════════════════════════════════════
# MWS-004  EULA ACCEPTANCE MODAL
# ═══════════════════════════════════════════════════════════════════════════

doc = Document()
set_margins(doc)

add_heading(doc, 'MWS-004 — EULA Acceptance Modal', level=1)
add_para(doc, 'Jira Story  ·  MillworkSuite UX Implementation', italic=True, color='445069', size=10)
add_divider(doc)
doc.add_paragraph()

add_meta_table(doc, [
    ('Story ID',      'MWS-004'),
    ('Epic',          'MWS — Public & Auth Pages'),
    ('Type',          'Story'),
    ('Priority',      'P0 — Blocker'),
    ('Story Points',  '5'),
    ('Labels',        'frontend, auth, legal, UX'),
    ('Trigger',       'Immediately after successful credential validation on Sign in'),
    ('Legal entity',  'VIEWPOINT DATA SYSTEMS, INC. d/b/a MillworkSuite'),
    ('EULA version',  'v1.0 · Effective 01 Jan 2026'),
])

# User story
add_heading(doc, 'User Story', level=2)
add_para(doc, (
    'As a first-time MillworkSuite user, I must read and accept the End User License Agreement '
    'before I can access the platform, so that Viewpoint Data Systems is legally protected and '
    'I understand my rights and obligations regarding the software.'
), italic=True, size=11)
doc.add_paragraph()

# ── Screenshots — all 4 states ──────────────────────────────────────────
add_heading(doc, 'Screen References — All States', level=2)

add_para(doc, 'STATE 1 — Initial (locked): Scroll hint visible, checkbox disabled, Accept button disabled', bold=True, size=10, color='1E5BA8')
add_screenshot(doc, f'{SS}/eula_locked.png',
               'Fig 1 — EULA modal, initial state. User has not yet scrolled. Checkbox and Accept button are locked.')

add_para(doc, 'STATE 2 — Scrolled to bottom: Scroll hint hidden, checkbox enabled, Accept button still disabled until checked', bold=True, size=10, color='1E5BA8')
add_screenshot(doc, f'{SS}/eula_unlocked.png',
               'Fig 2 — EULA modal after scrolling to bottom. Checkbox is now enabled. Accept button still disabled.')

add_para(doc, 'STATE 3 — Checkbox checked: Accept & Continue button becomes active', bold=True, size=10, color='1E5BA8')
add_screenshot(doc, f'{SS}/eula_accepted.png',
               'Fig 3 — Checkbox checked. "Accept & Continue →" button is now active and clickable.')

add_para(doc, 'STATE 4 — Decline confirmation dialog', bold=True, size=10, color='1E5BA8')
add_screenshot(doc, f'{SS}/eula_decline_confirm.png',
               'Fig 4 — Clicking "Decline" surfaces a confirmation dialog before signing the user out.')

add_divider(doc)

# ── Component breakdown ──────────────────────────────────────────────────
add_heading(doc, 'Component 1 — Modal Overlay', level=2)
bullets(doc, [
    '• Full-screen dark overlay (rgba backdrop) covers the entire app — nothing underneath is accessible.',
    '• role="dialog" aria-modal="true" aria-labelledby="eulaTitle" for screen-reader compliance.',
    '• Modal cannot be dismissed by clicking outside or pressing Escape — user must Accept or Decline.',
    '• Centered white modal card, max-width ~760 px, with three distinct zones: header, scrollable body, footer.',
])

add_heading(doc, 'Component 2 — Modal Header (dark navy bar)', level=2)
bullets(doc, [
    '• Background: dark navy gradient matching the app topbar (--blue-900 → --blue-800).',
    '• Left: MillworkSuite favicon/logo icon (28 × 28 px).',
    '• Title: "End User License Agreement" — white, bold, 15 px, display font.',
    '• Sub-title: "Please read the full agreement before continuing. You must scroll to the bottom to accept." — white/60% opacity, 11.5 px.',
    '• No close (×) button — intentional. User must choose Accept or Decline.',
])

add_heading(doc, 'Component 3 — Scrollable Body', level=2)
add_para(doc, 'The body is a fixed-height, overflow-y: scroll container with all legal content.', size=11)
add_para(doc, 'Content structure (14 numbered sections):', bold=True, size=11)
add_data_table(doc,
    ['Section', 'Title', 'Key Content'],
    [
        ['Preamble',   'Agreement parties',          'Identifies Viewpoint Data Systems, Inc. and the Customer'],
        ['Opening cap','Acceptance clause',           'All-caps: clicking Accept or using the software constitutes legal agreement'],
        ['§1',         'License Grant and Scope',    'Non-exclusive, non-transferable license; one license per user login'],
        ['§2',         'Restrictions',               '8 bullet prohibitions: no reverse-engineering, no sub-licensing, no competitive use, etc.'],
        ['§3',         'Fees',                        'Monthly billing in advance; non-refundable; USD only'],
        ['§4',         'Intellectual Property',       'Licensor retains all IP; Customer has no ownership interest'],
        ['§5',         'Use of Information / AI',    'AI Outputs are informational only; Customer assumes all risk; no warranty on AI outputs'],
        ['§6',         'Maintenance & Support',       'Updates provided per standard support program; Licensor may charge for support'],
        ['§7',         'Customer Data',               'Customer retains data ownership; Licensor may use data to improve AI models'],
        ['§8',         'Indemnification',             'Customer defends/indemnifies Licensor against third-party claims'],
        ['§9',         'Confidentiality',             'Mutual confidentiality obligations on proprietary information'],
        ['§10',        'Term and Termination',        'Licensor: 15-day cure or immediate; Customer: 30-day cure period'],
        ['§11',        'Limited Warranties',          'All-caps: Software provided "AS IS"; no implied warranties'],
        ['§12',        'Limitation of Liability',     'All-caps: No consequential/punitive damages; aggregate cap = total fees paid'],
        ['§13',        'Governing Law & Disputes',    'Florida law; Orange County courts; JURY TRIAL WAIVER (all-caps)'],
        ['§14',        'General Provisions',          'Export compliance, independent contractors, entire agreement, no assignment, severability, attorneys fees'],
        ['Closing cap','Final acceptance clause',     'All-caps restatement that accessing the software = binding agreement'],
    ]
)

add_para(doc, 'Typography rules within body:', bold=True, size=11)
bullets(doc, [
    '• H1: "END USER LICENSE AGREEMENT" — section title at top of body.',
    '• H2: Section numbers and titles (§1 through §14) — bold, with bottom border.',
    '• P: Body paragraphs — 13 px, line-height 1.6.',
    '• UL: Bulleted restriction list in §2.',
    '• .eula-caps: All-caps blocks (preamble caps, §11, §12, §13 jury waiver, closing caps) — '
      'monospace or uppercase font, smaller size (~11 px), background tint, used for legally significant passages.',
    '• .eula-preamble: Opening paragraph — slightly larger, 13 px, with bottom border.',
])

add_heading(doc, 'Component 4 — Scroll Gate', level=2)
bullets(doc, [
    '• A "Scroll down to read the full agreement ↓" hint strip is shown at the bottom of the body on load.',
    '• The hint is ONLY hidden once the user scrolls to within 140 px of the bottom.',
    '• Simultaneously, the checkbox is enabled (disabled attribute removed) and its label loses the "locked" style.',
    '• Logic is event-driven on the scroll event of the eulaBody element.',
    '• The user CANNOT check the checkbox or click Accept without having scrolled through the entire document.',
])

add_heading(doc, 'Component 5 — Footer', level=2)
add_para(doc, '5a. Checkbox + Agreement Label', bold=True, size=11, color='1E5BA8')
bullets(doc, [
    '• Checkbox input (type="checkbox", id="eulaAgree") — initially disabled.',
    '• Label text: "I have read the full agreement and I accept the MillworkSuite End User License Agreement on behalf of myself and my organization."',
    '• Label text is muted/grey ("locked" class) while checkbox is disabled; turns dark ink when enabled.',
    '• The checkbox row has a highlighted background when checked (--blue-50 tint, blue border).',
    '• Checking the checkbox does NOT yet trigger any navigation — it only enables the Accept button.',
])

add_para(doc, '5b. Version Tag', bold=True, size=11, color='1E5BA8')
bullets(doc, [
    '• Bottom-left text: "MillworkSuite EULA · v1.0 · Effective 01 Jan 2026 · Viewpoint Data Systems, Inc."',
    '• Font: 10.5 px, muted ink, left-aligned.',
])

add_para(doc, '5c. Action Buttons (right-aligned)', bold=True, size=11, color='1E5BA8')
add_data_table(doc,
    ['Button', 'Style', 'Enabled when', 'On click'],
    [
        ['Decline',           'Ghost / secondary',   'Always',                              'Opens Decline Confirmation dialog (see Component 6)'],
        ['Accept & Continue →', 'Primary blue + arrow', 'Checkbox is checked', 'Stores acceptance timestamp in localStorage (mws_eula_v1), closes modal, navigates to App Hub, shows success toast'],
    ]
)

add_heading(doc, 'Component 6 — Decline Confirmation Dialog', level=2)
add_para(doc, 'Clicking "Decline" does NOT immediately sign the user out. It first surfaces a confirmation dialog.', size=11)
bullets(doc, [
    '• Dialog title: "Decline agreement?"',
    '• Body: "You cannot access MillworkSuite without accepting the EULA. You will be signed out."',
    '• Style: danger (red icon/accent).',
    '• Two buttons: Cancel (returns to EULA modal, no action taken) and Confirm (signs user out → navigates to /login).',
    '• Dialog appears as a second modal layer on top of the EULA modal.',
])

add_divider(doc)

# ── State machine ────────────────────────────────────────────────────────
add_heading(doc, 'State Machine', level=2)
add_data_table(doc,
    ['State', 'Scroll hint', 'Checkbox', 'Accept button', 'Trigger to advance'],
    [
        ['1. Initial',         'Visible',  'Disabled (grey)',   'Disabled (grey)',  'User scrolls to within 140 px of bottom'],
        ['2. Scrolled',        'Hidden',   'Enabled (active)',  'Disabled (grey)',  'User checks the checkbox'],
        ['3. Ready to accept', 'Hidden',   'Checked (blue)',    'Enabled (blue)',   'User clicks "Accept & Continue"'],
        ['4. Accepted',        '—',        '—',                 '—',                'Modal closes; App Hub loads; success toast shown'],
        ['Decline path',       'Any',      'Any',               'Any',              'User clicks Decline → Confirm → returns to login'],
    ]
)

# ── Post-acceptance ──────────────────────────────────────────────────────
add_heading(doc, 'Post-Acceptance Behaviour', level=2)
bullets(doc, [
    '• localStorage key mws_eula_v1 is set to the ISO timestamp of acceptance (e.g. "2026-05-30T01:14:00.000Z").',
    '• On subsequent sign-ins: showEula() detects the key, clears it, and still shows the EULA — ensuring the modal always shows. (Current prototype behavior: always show for demo purposes.)',
    '• Production behavior: If mws_eula_v1 exists in the user\'s server-side profile, skip EULA and go directly to App Hub.',
    '• On Accept: success toast appears top-right — "Agreement accepted · Welcome to MillworkSuite, [Name]."',
    '• The acceptance timestamp and version (v1.0) must be stored server-side per user for compliance/audit.',
])

# ── Users table EULA status ──────────────────────────────────────────────
add_heading(doc, 'Admin Visibility — Users & Roles Screen', level=2)
add_para(doc, (
    'The Users & Roles admin screen shows a per-user EULA column with three possible states:'
), size=11)
add_data_table(doc,
    ['Status', 'Icon', 'Color', 'Meaning'],
    [
        ['Accepted',       'Checkmark ✓',   'Green',  'User has accepted — shows acceptance date (e.g. "19 May 26")'],
        ['Pending login',  'Info circle ℹ', 'Amber',  'User account exists but has not yet signed in for the first time'],
        ['Not accepted',   'Warning ⚠',     'Red',    'User signed in but declined or has not completed the flow'],
    ]
)
add_para(doc, 'Company admins can see EULA status for all users but cannot accept on their behalf.', italic=True, size=10, color='445069')
doc.add_paragraph()

# ── Accessibility ────────────────────────────────────────────────────────
add_heading(doc, 'Accessibility Requirements', level=2)
bullets(doc, [
    '• role="dialog" aria-modal="true" aria-labelledby="eulaTitle" on the overlay element.',
    '• Focus must be trapped inside the modal while it is open (tab cycles only through modal elements).',
    '• Escape key must NOT close the modal (user must explicitly choose Accept or Decline).',
    '• Checkbox must have an associated <label> element (not aria-label alone).',
    '• Accept button disabled state must use the disabled HTML attribute (not just CSS opacity).',
    '• Success toast after acceptance must be announced via aria-live="polite".',
    '• All-caps legal sections must remain readable by screen readers (use CSS text-transform, not HTML uppercase).',
    '• Scroll container must be keyboard-scrollable (focusable, responds to arrow keys/Page Down).',
])
doc.add_paragraph()

# ── ACs ──────────────────────────────────────────────────────────────────
add_divider(doc)
add_heading(doc, 'Acceptance Criteria', level=2)
add_ac_table(doc, [
    'EULA modal appears immediately after successful credential validation on first sign-in.',
    'Full-screen overlay prevents interaction with any app content behind the modal.',
    'Modal header shows MillworkSuite logo, title "End User License Agreement", and scroll instruction sub-title.',
    'No close (×) button is present — modal can only be dismissed via Accept or Decline.',
    'Body renders all 14 sections in correct order with proper heading hierarchy (H1, H2, P, UL).',
    'All-caps legal blocks (.eula-caps) render in distinct monospace/uppercase styling.',
    'Scroll hint ("Scroll down to read the full agreement ↓") is visible on load and hidden after scrolling.',
    'Checkbox is disabled and greyed-out until user scrolls within 140 px of the bottom.',
    'Checkbox label is muted ("locked" styling) while disabled; becomes dark ink when enabled.',
    'Accept & Continue button remains disabled until both: (a) user has scrolled to bottom AND (b) checkbox is checked.',
    'Clicking "Decline" opens the confirmation dialog with correct title, body text, and danger styling.',
    'Confirming decline: modal closes, user is navigated back to /login.',
    'Cancelling decline: confirmation dialog closes, user returns to EULA modal (state preserved).',
    'Clicking "Accept & Continue": modal closes, mws_eula_v1 timestamp stored, user navigates to App Hub.',
    'Success toast appears top-right: "Agreement accepted · Welcome to MillworkSuite, [Name]."',
    'Footer version string reads: "MillworkSuite EULA · v1.0 · Effective 01 Jan 2026 · Viewpoint Data Systems, Inc."',
    'Users & Roles admin table shows per-user EULA status (Accepted + date, Pending login, Not accepted).',
    'Focus is trapped inside the modal while it is open.',
    'Scroll container is keyboard-navigable (Page Down / arrow keys).',
    'Page passes WCAG 2.1 AA audit for the modal.',
])

doc.save(f'{OUT}/MWS-004-EULA-Acceptance.docx')
print('✓ MWS-004-EULA-Acceptance.docx')
